"""
Document Parser - Handles PDF/DOC requirement documents and prepares them for Llama synthesis
"""

import os
import json
from typing import List, Dict, Any, Optional
from dataclasses import dataclass
import PyPDF2
import pdfplumber
from docx import Document
from unstructured.partition.auto import partition
from langchain.text_splitter import RecursiveCharacterTextSplitter
import tiktoken

@dataclass
class RequirementTicket:
    """Structured representation of a requirement ticket"""
    page_number: int
    title: str
    description: str
    features: List[str]
    priority: str
    acceptance_criteria: List[str]
    technical_notes: str
    raw_text: str

@dataclass
class ParsedDocument:
    """Complete parsed document structure"""
    filename: str
    total_pages: int
    requirements: List[RequirementTicket]
    summary: str
    metadata: Dict[str, Any]

class DocumentParser:
    """Handles parsing of requirement documents (PDF, DOCX)"""
    
    def __init__(self):
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=1000,
            chunk_overlap=200,
            separators=["\n\n", "\n", ".", "!", "?", ",", " ", ""]
        )
        self.tokenizer = tiktoken.get_encoding("cl100k_base")
    
    def parse_document(self, file_path: str) -> ParsedDocument:
        """Parse document and extract requirement tickets"""
        
        file_extension = os.path.splitext(file_path)[1].lower()
        
        if file_extension == '.pdf':
            return self._parse_pdf(file_path)
        elif file_extension in ['.docx', '.doc']:
            return self._parse_docx(file_path)
        else:
            raise ValueError(f"Unsupported file type: {file_extension}")
    
    def _parse_pdf(self, file_path: str) -> ParsedDocument:
        """Parse PDF document using multiple strategies"""
        
        # Strategy 1: Use unstructured (Meta's library)
        try:
            return self._parse_with_unstructured(file_path)
        except Exception as e:
            print(f"Unstructured parsing failed: {e}")
        
        # Strategy 2: Fallback to pdfplumber
        try:
            return self._parse_with_pdfplumber(file_path)
        except Exception as e:
            print(f"Pdfplumber parsing failed: {e}")
        
        # Strategy 3: Fallback to PyPDF2
        return self._parse_with_pypdf2(file_path)
    
    def _parse_with_unstructured(self, file_path: str) -> ParsedDocument:
        """Parse using Meta's unstructured library"""
        
        elements = partition(filename=file_path)
        pages = self._group_elements_by_page(elements)
        
        requirements = []
        for page_num, page_elements in pages.items():
            page_text = "\n".join([str(elem) for elem in page_elements])
            requirement = self._extract_requirement_from_text(page_text, page_num)
            if requirement:
                requirements.append(requirement)
        
        return ParsedDocument(
            filename=os.path.basename(file_path),
            total_pages=len(pages),
            requirements=requirements,
            summary=self._generate_document_summary(requirements),
            metadata={"parser": "unstructured", "file_type": "pdf"}
        )
    
    def _parse_with_pdfplumber(self, file_path: str) -> ParsedDocument:
        """Parse using pdfplumber for better text extraction"""
        
        requirements = []
        with pdfplumber.open(file_path) as pdf:
            for page_num, page in enumerate(pdf.pages, 1):
                text = page.extract_text()
                if text and text.strip():
                    requirement = self._extract_requirement_from_text(text, page_num)
                    if requirement:
                        requirements.append(requirement)
        
        return ParsedDocument(
            filename=os.path.basename(file_path),
            total_pages=len(pdf.pages),
            requirements=requirements,
            summary=self._generate_document_summary(requirements),
            metadata={"parser": "pdfplumber", "file_type": "pdf"}
        )
    
    def _parse_with_pypdf2(self, file_path: str) -> ParsedDocument:
        """Fallback parsing with PyPDF2"""
        
        requirements = []
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            
            for page_num, page in enumerate(pdf_reader.pages, 1):
                text = page.extract_text()
                if text and text.strip():
                    requirement = self._extract_requirement_from_text(text, page_num)
                    if requirement:
                        requirements.append(requirement)
        
        return ParsedDocument(
            filename=os.path.basename(file_path),
            total_pages=len(pdf_reader.pages),
            requirements=requirements,
            summary=self._generate_document_summary(requirements),
            metadata={"parser": "pypdf2", "file_type": "pdf"}
        )
    
    def _parse_docx(self, file_path: str) -> ParsedDocument:
        """Parse DOCX document"""
        
        doc = Document(file_path)
        requirements = []
        
        # Group paragraphs by sections (assuming each section is a requirement)
        current_section = []
        section_num = 1
        
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if not text:
                continue
            
            # Check if this is a new section (heuristic)
            if self._is_new_section(paragraph):
                if current_section:
                    section_text = "\n".join(current_section)
                    requirement = self._extract_requirement_from_text(section_text, section_num)
                    if requirement:
                        requirements.append(requirement)
                    section_num += 1
                    current_section = []
            
            current_section.append(text)
        
        # Handle last section
        if current_section:
            section_text = "\n".join(current_section)
            requirement = self._extract_requirement_from_text(section_text, section_num)
            if requirement:
                requirements.append(requirement)
        
        return ParsedDocument(
            filename=os.path.basename(file_path),
            total_pages=section_num,
            requirements=requirements,
            summary=self._generate_document_summary(requirements),
            metadata={"parser": "python-docx", "file_type": "docx"}
        )
    
    def _extract_requirement_from_text(self, text: str, page_num: int) -> Optional[RequirementTicket]:
        """Extract structured requirement from text using heuristics and patterns"""
        
        # Clean and normalize text
        text = self._clean_text(text)
        
        if not text or len(text.strip()) < 50:  # Skip very short sections
            return None
        
        # Extract title (usually first line or after "Title:" pattern)
        title = self._extract_title(text)
        
        # Extract features (look for bullet points, numbered lists)
        features = self._extract_features(text)
        
        # Extract priority (look for priority indicators)
        priority = self._extract_priority(text)
        
        # Extract acceptance criteria
        acceptance_criteria = self._extract_acceptance_criteria(text)
        
        # Extract technical notes
        technical_notes = self._extract_technical_notes(text)
        
        # Generate description (remaining text)
        description = self._generate_description(text, title, features, acceptance_criteria)
        
        return RequirementTicket(
            page_number=page_num,
            title=title,
            description=description,
            features=features,
            priority=priority,
            acceptance_criteria=acceptance_criteria,
            technical_notes=technical_notes,
            raw_text=text
        )
    
    def _clean_text(self, text: str) -> str:
        """Clean and normalize text"""
        # Remove excessive whitespace
        text = " ".join(text.split())
        # Remove common PDF artifacts
        text = text.replace("", "")  # Form feed
        return text.strip()
    
    def _extract_title(self, text: str) -> str:
        """Extract title from text"""
        lines = text.split('\n')
        
        # Look for title patterns
        for line in lines[:3]:  # Check first 3 lines
            line = line.strip()
            if line and len(line) < 100 and not line.startswith(('•', '-', '1.', '2.')):
                return line
        
        # Fallback: use first meaningful line
        for line in lines:
            line = line.strip()
            if line and len(line) > 10 and len(line) < 200:
                return line
        
        return "Untitled Requirement"
    
    def _extract_features(self, text: str) -> List[str]:
        """Extract features from text"""
        features = []
        
        # Look for bullet points and numbered lists
        lines = text.split('\n')
        for line in lines:
            line = line.strip()
            if line.startswith(('•', '-', '*', '1.', '2.', '3.')):
                feature = line.lstrip('•-*1234567890. ').strip()
                if feature and len(feature) > 5:
                    features.append(feature)
        
        return features
    
    def _extract_priority(self, text: str) -> str:
        """Extract priority from text"""
        text_lower = text.lower()
        
        if any(word in text_lower for word in ['high priority', 'critical', 'urgent']):
            return "High"
        elif any(word in text_lower for word in ['medium priority', 'normal']):
            return "Medium"
        elif any(word in text_lower for word in ['low priority', 'nice to have']):
            return "Low"
        
        return "Medium"  # Default
    
    def _extract_acceptance_criteria(self, text: str) -> List[str]:
        """Extract acceptance criteria from text"""
        criteria = []
        
        # Look for acceptance criteria patterns
        lines = text.split('\n')
        in_criteria_section = False
        
        for line in lines:
            line = line.strip()
            if any(phrase in line.lower() for phrase in ['acceptance criteria', 'acceptance test', 'criteria:']):
                in_criteria_section = True
                continue
            
            if in_criteria_section and line:
                if line.startswith(('•', '-', '*', '1.', '2.')):
                    criteria.append(line.lstrip('•-*1234567890. ').strip())
                elif line and len(line) > 10:
                    criteria.append(line)
        
        return criteria
    
    def _extract_technical_notes(self, text: str) -> str:
        """Extract technical notes from text"""
        # Look for technical sections
        lines = text.split('\n')
        technical_lines = []
        in_technical_section = False
        
        for line in lines:
            line = line.strip()
            if any(phrase in line.lower() for phrase in ['technical', 'implementation', 'notes:', 'tech:']):
                in_technical_section = True
                continue
            
            if in_technical_section and line:
                technical_lines.append(line)
        
        return " ".join(technical_lines)
    
    def _generate_description(self, text: str, title: str, features: List[str], criteria: List[str]) -> str:
        """Generate description from remaining text"""
        # Remove already extracted parts
        description = text
        
        # Remove title
        if title in description:
            description = description.replace(title, "")
        
        # Remove features
        for feature in features:
            if feature in description:
                description = description.replace(feature, "")
        
        # Remove criteria
        for criterion in criteria:
            if criterion in description:
                description = description.replace(criterion, "")
        
        # Clean up
        description = " ".join(description.split())
        return description[:500]  # Limit length
    
    def _is_new_section(self, paragraph) -> bool:
        """Check if paragraph starts a new section"""
        text = paragraph.text.strip()
        
        # Check for heading styles
        if paragraph.style.name.startswith('Heading'):
            return True
        
        # Check for common section patterns
        if any(pattern in text.lower() for pattern in ['requirement', 'feature', 'story', 'ticket']):
            return True
        
        # Check for numbering patterns
        if text and text[0].isdigit() and '.' in text[:5]:
            return True
        
        return False
    
    def _group_elements_by_page(self, elements) -> Dict[int, List]:
        """Group unstructured elements by page"""
        pages = {}
        current_page = 1
        
        for element in elements:
            # Try to extract page number from element metadata
            if hasattr(element, 'metadata') and 'page_number' in element.metadata:
                current_page = element.metadata['page_number']
            
            if current_page not in pages:
                pages[current_page] = []
            
            pages[current_page].append(element)
        
        return pages
    
    def _generate_document_summary(self, requirements: List[RequirementTicket]) -> str:
        """Generate summary of the document"""
        if not requirements:
            return "No requirements found in document."
        
        summary_parts = [
            f"Document contains {len(requirements)} requirement tickets:",
            f"- High priority: {len([r for r in requirements if r.priority == 'High'])}",
            f"- Medium priority: {len([r for r in requirements if r.priority == 'Medium'])}",
            f"- Low priority: {len([r for r in requirements if r.priority == 'Low'])}",
            f"- Total features: {sum(len(r.features) for r in requirements)}"
        ]
        
        return " ".join(summary_parts)
    
    def chunk_for_llama(self, parsed_doc: ParsedDocument, max_tokens: int = 4000) -> List[str]:
        """Chunk document for Llama processing"""
        chunks = []
        
        # Create comprehensive text representation
        full_text = f"Document: {parsed_doc.filename}\n"
        full_text += f"Summary: {parsed_doc.summary}\n\n"
        
        for req in parsed_doc.requirements:
            req_text = f"Requirement {req.page_number}:\n"
            req_text += f"Title: {req.title}\n"
            req_text += f"Priority: {req.priority}\n"
            req_text += f"Description: {req.description}\n"
            req_text += f"Features: {', '.join(req.features)}\n"
            req_text += f"Acceptance Criteria: {', '.join(req.acceptance_criteria)}\n"
            if req.technical_notes:
                req_text += f"Technical Notes: {req.technical_notes}\n"
            req_text += "\n"
            
            full_text += req_text
        
        # Split into chunks
        text_chunks = self.text_splitter.split_text(full_text)
        
        # Ensure chunks don't exceed token limit
        for chunk in text_chunks:
            tokens = len(self.tokenizer.encode(chunk))
            if tokens <= max_tokens:
                chunks.append(chunk)
            else:
                # Further split if needed
                sub_chunks = self.text_splitter.split_text(chunk)
                chunks.extend(sub_chunks)
        
        return chunks 